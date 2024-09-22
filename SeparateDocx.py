from docx import Document
import googletrans
translator = googletrans.Translator()
def read_paragraphs_from_docx(file_path, min_length=30):
    """
    파일에서 문단을 읽고 30자 이하인 문단을 제외하는 함수.
    """
    # 문서 열기
    doc = Document(file_path)
    # 문단을 리스트로 추출 및 필터링 (문단의 길이가 min_length 이상인 경우에만 포함)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() and len(para.text.strip()) > min_length]
    return paragraphs

def save_paragraphs_to_docx(paragraphs, output_file_path):
    """
    필터링된 문단을 새 .docx 파일로 저장하는 함수.
    """
    # 새 문서 생성
    new_doc = Document()
    # 문단을 새 문서에 추가
    for para in paragraphs:
        result = translator.translate(para, dest="ko", src="en")
        new_doc.add_paragraph(result.text)
    # 파일로 저장
    new_doc.save(output_file_path)

# 사용 예시
input_docx_path = r'C:\Users\JHJH\Desktop\PyProject\End-to-End Object Detection with Transformers.docx'  # 읽을 파일 경로
output_docx_path = 'output_file.docx'  # 저장할 파일 경로

# 파일 읽기 (30자 이하 문단 제외)
paragraphs = read_paragraphs_from_docx(input_docx_path, min_length=150)
# 파일 저장
save_paragraphs_to_docx(paragraphs, output_docx_path)

print("새로운 .docx 파일에 저장되었습니다.")
